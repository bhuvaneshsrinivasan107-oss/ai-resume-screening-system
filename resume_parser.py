import os
import re
import io
import shutil
import logging
import tempfile

import pdfplumber
import pytesseract

from PIL import Image
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger("resume_parser")


# ============================================================
# TESSERACT CONFIGURATION (cross-platform)
# ============================================================

def _find_tesseract():
    if os.name == "nt":
        win_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.isfile(win_path):
            return win_path
    unix_path = shutil.which("tesseract")
    if unix_path:
        return unix_path
    return None

TESSERACT_PATH = _find_tesseract()

if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    logger.info("Tesseract found at: %s", TESSERACT_PATH)
else:
    logger.warning("Tesseract binary not found on this system. OCR will be unavailable.")


# ============================================================
# CLEAN TEXT
# ============================================================

def clean_text(text):

    if not text:
        return ""

    text = text.replace("\x00", " ")

    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# ============================================================
# EXTRACT TEXT FROM NORMAL PDF
# ============================================================

def _extract_with_pymupdf(file_bytes):
    """Extract text using PyMuPDF (fitz) - fastest and most reliable."""

    try:

        # Prefer the modern `pymupdf` import; fall back to the
        # classic `fitz` name for older versions of the package.
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz

        doc = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        pages_text = []

        for page in doc:

            page_text = page.get_text()

            if page_text:

                pages_text.append(
                    page_text
                )

        doc.close()

        if pages_text:

            return "\n".join(pages_text)

    except Exception as e:

        logger.warning("PyMuPDF extraction failed: %s", e)

    return ""


def _extract_with_pdfplumber(file_bytes):
    """Extract text using pdfplumber (fallback)."""

    try:

        with pdfplumber.open(
            io.BytesIO(file_bytes)
        ) as pdf:

            pages_text = []

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    pages_text.append(
                        page_text
                    )

            return "\n".join(pages_text)

    except Exception as e:

        logger.warning("pdfplumber extraction failed: %s", e)

    return ""


def _extract_with_pypdf(file_bytes):
    """Extract text using pypdf (final fallback)."""

    try:

        reader = PdfReader(
            io.BytesIO(file_bytes)
        )

        pages_text = []

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:

                pages_text.append(
                    page_text
                )

        return "\n".join(pages_text)

    except Exception as e:

        logger.warning("pypdf extraction failed: %s", e)

    return ""


def extract_pdf_text(file_bytes):

    text = ""

    # --------------------------------------------------------
    # Method 1: PyMuPDF (fitz) - primary
    # --------------------------------------------------------

    text = _extract_with_pymupdf(file_bytes)

    if text.strip():

        return clean_text(text)

    # --------------------------------------------------------
    # Method 2: pdfplumber
    # --------------------------------------------------------

    text = _extract_with_pdfplumber(file_bytes)

    if text.strip():

        return clean_text(text)

    # --------------------------------------------------------
    # Method 3: pypdf
    # --------------------------------------------------------

    text = _extract_with_pypdf(file_bytes)

    return clean_text(text)


# ============================================================
# TESSERACT VALIDATION
# ============================================================

def tesseract_available():
    """Return True if the Tesseract OCR engine is usable.

    This is a development-side check: it verifies the binary is
    reachable. On Render the Dockerfile installs tesseract-ocr,
    so this should return True inside the container.
    """
    if not TESSERACT_PATH:
        return False
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


# ============================================================
# OCR IMAGE PREPROCESSING
# ============================================================

def _ocrize(pil_image):
    """Preprocess a page image to improve OCR accuracy.

    Converts to grayscale and boosts contrast so Tesseract
    can read scanned resumes (which are often light with
    normal white backgrounds) more reliably.
    """
    gray = pil_image.convert("L")
    try:
        from PIL import ImageEnhance
        gray = ImageEnhance.Contrast(gray).enhance(2.0)
    except Exception:
        pass
    return gray


def _ocr_image_best(pil_image, config=""):
    """Run OCR on an image, trying multiple rotations if needed.

    Some scanned PDFs have the page rotated 90/180/270 degrees.
    We OCR the original plus each rotation and keep the result
    that contains the most text.
    """
    best_text = ""
    best_count = 0

    candidates = [0, 180, 90, 270]

    for angle in candidates:

        try:

            img = pil_image

            if angle != 0:

                img = pil_image.rotate(
                    angle,
                    expand=True
                )

            processed = _ocrize(img)

            text = pytesseract.image_to_string(
                processed,
                config=config
            )

            if text:

                cleaned = clean_text(text)

                # Non-blank lines roughly indicate useful content.
                count = len(
                    [
                        line
                        for line in cleaned.splitlines()
                        if line.strip()
                    ]
                )

                if count > best_count:

                    best_count = count

                    best_text = cleaned

        except Exception as e:

            logger.warning(
                "OCR rotation %s failed: %s",
                angle,
                e,
            )

    return best_text


# ============================================================
# OCR SCANNED PDF
# ============================================================

def extract_pdf_with_ocr(file_bytes):

    if not TESSERACT_PATH:
        logger.warning("OCR skipped: tesseract binary not available.")
        return ""

    # --------------------------------------------------------
    # Method 1: pdf2image + Tesseract (uses Poppler to render
    # each PDF page to a Pillow image, then OCRs it). This is
    # the most reliable general-purpose OCR path for scanned
    # and image-based PDFs.
    # --------------------------------------------------------

    tmp_dir = None

    try:

        from pdf2image import convert_from_bytes

        # Render with a temp dir so `pdftoppm` does not leak
        # any files into the repo or working directory.
        tmp_dir = tempfile.mkdtemp(
            prefix="resume_pdf_ocr_"
        )

        images = convert_from_bytes(
            file_bytes,
            dpi=300,
            output_folder=tmp_dir
        )

        extracted_pages = []

        for image in images:

            text = _ocr_image_best(
                image
            )

            if text:

                extracted_pages.append(
                    text
                )

        if extracted_pages:

            return clean_text(
                "\n".join(
                    extracted_pages
                )
            )

    except Exception as e:

        logger.warning("pdf2image OCR method failed: %s", e)

    finally:

        if tmp_dir and os.path.isdir(tmp_dir):

            shutil.rmtree(tmp_dir, ignore_errors=True)

    # --------------------------------------------------------
    # Method 2: PyMuPDF page rendering + Tesseract (no Poppler)
    #
    # If Poppler is unavailable, render each page directly to
    # a high-resolution image with PyMuPDF and OCR it.
    # --------------------------------------------------------

    try:

        try:
            import pymupdf as fitz
        except ImportError:
            import fitz

        doc = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        extracted_pages = []

        for page in doc:

            # Render the page to a high-resolution pixmap
            # (300 DPI gives Tesseract more detail to work with).
            page_image = page.get_pixmap(
                dpi=300
            )

            pil_image = Image.open(
                io.BytesIO(
                    page_image.tobytes(
                        "png"
                    )
                )
            )

            text = _ocr_image_best(
                pil_image
            )

            if text and text.strip():

                extracted_pages.append(
                    text
                )

        doc.close()

        if extracted_pages:

            return clean_text(
                "\n".join(
                    extracted_pages
                )
            )

    except Exception as e:

        logger.warning("PyMuPDF OCR method failed: %s", e)

    return ""


# ============================================================
# EXTRACT DOCX TEXT
# ============================================================

def extract_docx_text(file_bytes):

    try:

        document = Document(
            io.BytesIO(file_bytes)
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                paragraphs.append(
                    paragraph.text
                )

        # Also extract tables
        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    row_text.append(
                        cell.text
                    )

                paragraphs.append(
                    " ".join(row_text)
                )

        return clean_text(
            "\n".join(paragraphs)
        )

    except Exception as e:

        logger.error(
            "DOCX ERROR: %s",
            e,
        )

        return ""


# ============================================================
# MAIN TEXT EXTRACTION
# ============================================================

def extract_text(uploaded_file):
    """Compatibility wrapper: return extracted text as a string.

    Returns "" (empty) when no text can be extracted. Prefer
    `extract_resume_with_result` when a per-file error reason is
    needed for the UI.
    """
    text, _ = extract_resume_with_result(uploaded_file)
    return text


def extract_resume_with_result(uploaded_file):
    """Extract resume text and return (text, reason).

    reason is a dict with keys: 'filename', 'ok', 'message',
    'suggested_action'. This is used by the UI to report failures
    per file without crashing the whole screening run.
    """

    filename = getattr(
        uploaded_file,
        'name',
        'unknown'
    )

    def _fail(reason, action):
        return "", {
            "filename": filename,
            "ok": False,
            "message": reason,
            "suggested_action": action,
        }

    try:

        file_name = filename.lower()

        file_bytes = uploaded_file.getvalue()

        if not file_bytes:

            return _fail(
                "The file is empty.",
                "Please upload a non-empty resume.",
            )

        # ----------------------------------------------------
        # PDF
        # ----------------------------------------------------

        if file_name.endswith(".pdf"):

            # First try normal PDF text extraction.
            text = extract_pdf_text(
                file_bytes
            )

            if text and len(text.strip()) >= 30:

                return text, {
                    "filename": filename,
                    "ok": True,
                    "message": "",
                    "suggested_action": "",
                }

            # Otherwise attempt OCR.
            logger.info(
                "Normal PDF extraction yielded insufficient text "
                "for '%s'. Trying OCR...",
                filename,
            )

            ocr_text = extract_pdf_with_ocr(
                file_bytes
            )

            if ocr_text and ocr_text.strip():

                return ocr_text, {
                    "filename": filename,
                    "ok": True,
                    "message": "",
                    "suggested_action": "",
                }

            # Determine a helpful reason based on environment.
            if not tesseract_available():

                return _fail(
                    "OCR engine (Tesseract) is not available on "
                    "the server, so this image-based resume could "
                    "not be read.",
                    "The server needs tesseract-ocr installed. "
                    "The Docker deployment installs it automatically.",
                )

            return _fail(
                "The PDF appears to contain no readable text and "
                "OCR could not extract any either.",
                "The PDF may be corrupted or contain an unsupported "
                "image format. Try a clearer scan.",
            )

        # ----------------------------------------------------
        # DOCX
        # ----------------------------------------------------

        elif file_name.endswith(".docx"):

            docx_text = extract_docx_text(
                file_bytes
            )

            if docx_text:

                return docx_text, {
                    "filename": filename,
                    "ok": True,
                    "message": "",
                    "suggested_action": "",
                }

            return _fail(
                "This DOCX document did not yield any readable text.",
                "The document may be empty or password-protected.",
            )

        # ----------------------------------------------------
        # TXT
        # ----------------------------------------------------

        elif file_name.endswith(".txt"):

            try:

                txt_text = clean_text(
                    file_bytes.decode(
                        "utf-8",
                        errors="ignore"
                    )
                )

                if txt_text:

                    return txt_text, {
                        "filename": filename,
                        "ok": True,
                        "message": "",
                        "suggested_action": "",
                    }

            except Exception:

                pass

            return _fail(
                "This text file yielded no readable content.",
                "The file may be empty.",
            )

        # ----------------------------------------------------
        # Unsupported format
        # ----------------------------------------------------

        return _fail(
            "Unsupported file format.",
            "Please upload a PDF or DOCX resume.",
        )

    except Exception as e:

        logger.error(
            "TEXT EXTRACTION ERROR for '%s': %s",
            filename,
            e,
        )

        return _fail(
            "The resume could not be read due to an unexpected "
            "processing error.",
            "The file may be corrupted. Please try another resume.",
        )


# ============================================================
# EXTRACT EMAIL
# ============================================================

def extract_email(text):

    pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"

    match = re.search(
        pattern,
        text
    )

    if match:

        email = match.group(0)

        # Remove OCR noise prefix like "W_" before the
        # real email address (e.g. "W_arun-kumar@email.com")
        email = re.sub(
            r"^[A-Za-z0-9]_(?=[A-Za-z0-9])",
            "",
            email
        )

        email = re.sub(
            r"^[^A-Za-z0-9]+",
            "",
            email
        )

        if email:

            return email

    return "Not Found"


# ============================================================
# EXTRACT PHONE
# ============================================================

def extract_phone(text):

    patterns = [

        r"\+91[\s-]?\d{10}",

        r"\b\d{10}\b",

        r"\+91[\s-]?\d{5}[\s-]?\d{5}"

    ]

    for pattern in patterns:

        match = re.search(
            pattern,
            text
        )

        if match:

            return match.group(0)

    return "Not Found"


# ============================================================
# EXTRACT NAME
# ============================================================

def extract_name(text, filename=""):

    lines = [

        line.strip()

        for line in text.splitlines()

        if line.strip()
    ]

    # --------------------------------------------------------
    # Try first few lines
    # --------------------------------------------------------

    for line in lines[:15]:

        clean_line = re.sub(
            r"[^A-Za-z .'-]",
            "",
            line
        ).strip()

        words = clean_line.split()

        if 2 <= len(words) <= 4:

            lower_line = clean_line.lower()

            blocked = [

                "resume",
                "curriculum",
                "vitae",
                "email",
                "phone",
                "mobile",
                "address",
                "profile",
                "objective",
                "summary",
                "skills",
                "education",
                "experience",
                "contact"
            ]

            if not any(
                word in lower_line
                for word in blocked
            ):

                if all(
                    re.match(
                        r"^[A-Za-z][A-Za-z.'-]*$",
                        word
                    )
                    for word in words
                ):

                    return clean_line

    # --------------------------------------------------------
    # Try filename
    # --------------------------------------------------------

    if filename:

        name = os.path.splitext(
            os.path.basename(filename)
        )[0]

        name = re.sub(
            r"resume",
            "",
            name,
            flags=re.IGNORECASE
        )

        name = re.sub(
            r"[_\-]+",
            " ",
            name
        )

        name = name.strip()

        if name:

            return name.title()

    return "Unknown Candidate"


# ============================================================
# SKILL EXTRACTION
# ============================================================

def extract_skills(text):

    skill_dictionary = [

        "python",
        "java",
        "c++",
        "c#",
        "c",

        "sql",
        "mysql",
        "postgresql",
        "mongodb",

        "excel",
        "power bi",
        "tableau",

        "pandas",
        "numpy",
        "scikit-learn",
        "sklearn",

        "machine learning",
        "deep learning",
        "artificial intelligence",

        "data science",
        "data analysis",
        "data analytics",

        "nlp",
        "natural language processing",

        "computer vision",

        "tensorflow",
        "pytorch",

        "aws",
        "azure",
        "google cloud",

        "docker",

        "git",
        "github",

        "flask",
        "fastapi",
        "streamlit",

        "html",
        "css",
        "javascript",
        "react",
        "node.js",

        "matplotlib",
        "seaborn",

        "statistics",
        "statistical analysis",

        "powerpoint",
        "ms office"
    ]

    text_lower = text.lower()

    found = []

    for skill in skill_dictionary:

        # Short skills (e.g. "c", "c#") must appear as
        # standalone words - otherwise "c" matches inside
        # any word containing the letter (e.g. "excel").
        if len(skill) <= 2:

            pattern = (
                r"(?<![A-Za-z0-9])"
                + re.escape(skill)
                + r"(?![A-Za-z0-9])"
            )

            present = (
                re.search(
                    pattern,
                    text_lower
                )
                is not None
            )

        else:

            present = skill in text_lower

        if present:

            if skill not in found:

                found.append(
                    skill
                )

    return found


# ============================================================
# EXTRACT CANDIDATE DETAILS
# ============================================================

def extract_candidate_details(
    text,
    filename=""
):

    text = clean_text(
        text
    )

    name = extract_name(
        text,
        filename
    )

    email = extract_email(
        text
    )

    phone = extract_phone(
        text
    )

    skills = extract_skills(
        text
    )

    return {

        "name": name,

        "email": email,

        "phone": phone,

        "skills": skills,

        "resume_text": text

    }


# ============================================================
# TEST
# ============================================================

if __name__ == "__main__":

    print(
        "Resume parser loaded successfully."
    )

    print(
        "Tesseract:",
        TESSERACT_PATH or "not found"
    )